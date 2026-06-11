#!/usr/bin/env python3
"""docx 题库解析器: 支持 "题目+【答案】+【解析】" 混排格式(答案文档)。
可单独 CLI 使用: python3 import_docx.py 答案.docx [更多.docx...] -o data/bank.json
"""
import re, json, sys, hashlib

Q_RE = re.compile(r'^(\d{1,3})[.、．]\s*(.+)')
OPT_RE = re.compile(r'^([A-H])[.、．]?\s*(.+)')
ANS_RE = re.compile(r'^(?:(\d{1,3})[.、．]?)?\s*【答案】\s*([A-H]+)')
EXP_RE = re.compile(r'^(?:(\d{1,3})[.、．]?)?\s*【解析】\s*(.*)')
PART_RE = re.compile(r'^[（(]\d{1,2}[）)]\s*(.+)')   # 问答题小问 （1）...
REF_RE = re.compile(r'^答[：:]\s*(.*)')              # 问答题参考答案(可选)


def norm_key(text):
    """题干去空白/标点差异后做去重键"""
    t = re.sub(r'[\s（）()　。.，,]', '', text)
    return hashlib.md5(t.encode()).hexdigest()[:12]


def _flush(questions, cur):
    """收集所有题目; 选择题无答案先保留, 末尾一起补答案"""
    if not cur:
        return
    if cur['type'] == 'choice' and cur['options']:
        questions.append(cur)
    elif cur['type'] == 'qa' and cur['stem']:
        questions.append(cur)


def parse_paragraphs(paras, source=''):
    questions, cur = [], None
    for raw in paras:
        line = raw.strip()
        if not line:
            continue
        m = ANS_RE.match(line)
        if m:
            qnum = m.group(1)
            # 编号答案块(末尾集中): 先flush待处理题; 内联答案(无编号): 保持cur
            if cur and qnum and cur not in questions:
                _flush(questions, cur)
                cur = None
            if qnum:
                i = int(qnum) - 1
                if 0 <= i < len(questions):
                    cur = questions[i]
            if cur:
                cur['answer'] = sorted(set(m.group(2)))
            continue
        m = EXP_RE.match(line)
        if m:
            qnum = m.group(1)
            if cur and qnum and cur not in questions:
                _flush(questions, cur)
                cur = None
            if qnum:
                i = int(qnum) - 1
                if 0 <= i < len(questions):
                    cur = questions[i]
            if cur:
                cur['explanation'] = (cur.get('explanation', '') + m.group(2)).strip()
            continue
        m = REF_RE.match(line)
        if m and cur and cur['type'] == 'qa':
            ans = m.group(1).strip()
            # 把"答："对齐到对应小问: 第N个答案配第N个小问, 供逐小问作答
            if cur['parts'] and len(cur['answers']) < len(cur['parts']):
                cur['answers'].append(ans)
            elif cur['answers']:
                cur['answers'][-1] = (cur['answers'][-1] + ans).strip()
            else:
                cur['answers'].append(ans)   # 无小问的整题答案
            cur['ref'] = (cur.get('ref', '') + ans).strip()
            continue
        m = OPT_RE.match(line)
        if m and cur and cur['type'] == 'choice':
            cur['options'][m.group(1)] = m.group(2).strip()
            continue
        m = PART_RE.match(line)
        if m and cur and cur['type'] == 'qa':
            # 子问题与答案交替出现; 答句内的编号如"（2）进行xxx"不含"？"则算续行
            if not cur.get('ref') or '？' in line or '?' in line:
                cur['parts'].append(line)
            continue
        m = Q_RE.match(line)
        # 题干以数字开头(如标题"6.11五百问")不是新题, 防误切
        if m and not m.group(2)[:1].isdigit():
            _flush(questions, cur)
            stem = m.group(2).strip()
            qa = bool(re.match(r'^问[：:]', stem))
            if qa:
                stem = re.sub(r'^问[：:]\s*', '', stem)
            cur = {'type': 'qa' if qa else 'choice', 'stem': stem, 'options': {},
                   'answer': None, 'explanation': '', 'parts': [], 'answers': [],
                   'ref': '', 'source': source}
            continue
        # 题干/解析的折行续行
        if cur:
            if cur['type'] == 'qa':
                if cur.get('ref'):
                    cur['ref'] += line
                    if cur['answers']:
                        cur['answers'][-1] += line
                elif cur['parts']:
                    cur['parts'][-1] += line
                else:
                    cur['stem'] += line
            elif cur.get('answer') is not None:
                cur['explanation'] = (cur.get('explanation', '') + line).strip()
            elif not cur['options']:
                cur['stem'] += line
    _flush(questions, cur) if cur not in questions else None
    # 过滤掉最终仍无答案的选择题
    questions = [q for q in questions if not (q['type'] == 'choice' and q['answer'] is None)]
    for q in questions:
        if q['type'] == 'choice':
            q['multi'] = len(q['answer']) > 1
            q.pop('parts', None); q.pop('ref', None); q.pop('answers', None)
        else:
            q['answer'] = []
            q['multi'] = False
        q['id'] = norm_key(q['stem'])
    return questions


def parse_docx(path):
    import docx
    d = docx.Document(path)
    paras = [p.text for p in d.paragraphs]
    # 表格里的文字也收(有些资料用表格排版)
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                paras.extend(p.text for p in cell.paragraphs)
    import os
    return parse_paragraphs(paras, source=os.path.basename(path))


def _natkey(s):
    """文件名自然排序键: 让 1,2,...,10 正确排序而非 1,10,2"""
    return [int(t) if t.isdigit() else t.lower()
            for t in re.split(r'(\d+)', s or '')]


def sort_bank_by_source(bank):
    """题库按 来源文件名(自然序) → 文件内原始顺序 稳定排序。
    原始顺序用每题首次出现的位置保留。"""
    for i, q in enumerate(bank):
        q.setdefault('_ord', i)
    bank.sort(key=lambda q: (_natkey(q.get('source', '')), q.get('_ord', 0)))
    for q in bank:
        q.pop('_ord', None)
    return bank


def merge_into_bank(bank, new_qs):
    """按题干去重合并; 新题覆盖同 id 旧题(资料更新场景)。返回(新增,更新)
    合并后整库按文件名自然序排列, 修复批量导入乱序问题。"""
    idx = {q['id']: i for i, q in enumerate(bank)}
    added = updated = 0
    for q in new_qs:
        if q['id'] in idx:
            bank[idx[q['id']]] = q
            updated += 1
        else:
            bank.append(q)
            added += 1
    sort_bank_by_source(bank)
    return added, updated


if __name__ == '__main__':
    args = [a for a in sys.argv[1:] if a != '-o']
    out = 'data/bank.json'
    if '-o' in sys.argv:
        out = sys.argv[sys.argv.index('-o') + 1]
        args.remove(out)
    try:
        bank = json.load(open(out, encoding='utf-8'))
    except Exception:
        bank = []
    for f in args:
        qs = parse_docx(f)
        a, u = merge_into_bank(bank, qs)
        print(f'{f}: 解析 {len(qs)} 题, 新增 {a}, 更新 {u}')
    json.dump(bank, open(out, 'w', encoding='utf-8'), ensure_ascii=False, indent=1)
    print(f'题库共 {len(bank)} 题 -> {out}')
